#!/usr/bin/env python3
"""
Convert Markdown articles to Word documents with image insertion and table support.
Supports headings, bold text, lists, images, tables, and basic formatting.
"""

import sys
import argparse
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_chinese_font(run, font_name='Microsoft YaHei', font_size=12, bold=False):
    """Set Chinese font for a run (default: Microsoft YaHei)."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def parse_table_line(line):
    """Parse a table row from markdown."""
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    
    cells = [cell.strip() for cell in line.split('|')]
    return cells


def is_table_separator(line):
    """Check if line is a table separator."""
    stripped = line.strip()
    if not stripped.startswith('|') or not stripped.endswith('|'):
        return False
    
    content = stripped[1:-1]
    return all(c in '-|: \t' for c in content)


def remove_image_description_section(md_content):
    """
    Remove image description section from the end of article.
    Matches patterns like:
    - **配图说明：** or **配图说明：**
    - ---
    - [[IMG: ...]] descriptions
    - 1. [[IMG: ...]] - description
    """
    lines = md_content.split('\n')
    
    # Find the start of image description section
    image_section_start = -1
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        # Match various image description headers
        if (stripped.startswith('**配图说明') or 
            stripped.startswith('**配图清单') or
            stripped.startswith('---') and i > len(lines) * 0.7):  # Horizontal rule near end
            # Check if next lines contain IMG placeholders
            for j in range(i+1, min(i+10, len(lines))):
                if '[[IMG:' in lines[j] or '配图' in lines[j]:
                    image_section_start = i
                    break
            if image_section_start >= 0:
                break
        # Also detect standalone IMG list at end
        if re.match(r'^\d+\.[\s\t]*\[\[IMG:', stripped) and i > len(lines) * 0.7:
            # Look backwards for the section start
            for k in range(i-1, max(0, i-10), -1):
                if lines[k].strip().startswith('**') or lines[k].strip().startswith('---'):
                    image_section_start = k
                    break
            if image_section_start < 0:
                image_section_start = i
            break
    
    # If found, remove from that point onwards
    if image_section_start >= 0:
        # Keep content before the image section
        return '\n'.join(lines[:image_section_start])
    
    return md_content


def parse_markdown(md_content):
    """Parse markdown content into structured elements."""
    # Remove image description section first
    md_content = remove_image_description_section(md_content)
    
    elements = []
    lines = md_content.split('\n')
    
    i = 0
    first_heading = True
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Skip disclaimer footer
        if line.startswith('*本文基于') or line.startswith('*本文仅供'):
            i += 1
            continue
        
        # Skip table of contents
        if line.startswith('| 编号') or line.startswith('|------'):
            i += 1
            continue
        if re.match(r'\|\s*\d+\s*\|', line):
            i += 1
            continue
        
        # Table detection
        if line.startswith('|'):
            table_rows = []
            is_table = False
            j = i
            
            while j < len(lines) and lines[j].strip().startswith('|'):
                if is_table_separator(lines[j]):
                    is_table = True
                table_rows.append(lines[j].strip())
                j += 1
            
            if len(table_rows) >= 2 and is_table:
                headers = None
                rows = []
                
                for row_line in table_rows:
                    if is_table_separator(row_line):
                        continue
                    cells = parse_table_line(row_line)
                    if headers is None:
                        headers = cells
                    else:
                        rows.append(cells)
                
                elements.append({
                    'type': 'table',
                    'headers': headers,
                    'rows': rows
                })
                i = j
                continue
        
        # Image placeholder (support both [[IMG:...]] and **[[IMG:...]]**)
        img_match = re.search(r'\[\[IMG:\s*(.*?)\]\]', line)
        if img_match:
            elements.append({
                'type': 'image_placeholder',
                'description': img_match.group(1)
            })
            i += 1
            continue
        
        # Heading 1 - keep first one (article title)
        if line.startswith('# ') and not line.startswith('## '):
            elements.append({
                'type': 'heading1',
                'content': line[2:].strip()
            })
            i += 1
            continue
        
        # Heading 2
        if line.startswith('## ') and not line.startswith('### '):
            elements.append({
                'type': 'heading2',
                'content': line[3:].strip()
            })
            i += 1
            continue
        
        # Heading 3
        if line.startswith('### '):
            elements.append({
                'type': 'heading3',
                'content': line[4:].strip()
            })
            i += 1
            continue
        
        # Horizontal rule
        if line == '---':
            elements.append({'type': 'hr'})
            i += 1
            continue
        
        # Editor/Reviewer lines (内容编辑丨xxx, 内容审核丨xxx)
        if line.startswith('内容编辑丨') or line.startswith('内容审核丨'):
            elements.append({
                'type': 'editor_line',
                'content': line
            })
            i += 1
            continue
        
        # List item
        if line.startswith('- ') or line.startswith('* '):
            list_items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                list_items.append(lines[i].strip()[2:])
                i += 1
            elements.append({
                'type': 'list',
                'items': list_items
            })
            continue
        
        # Regular paragraph
        elements.append({
            'type': 'paragraph',
            'content': line
        })
        i += 1
    
    return elements


def process_inline_formatting(paragraph, text):
    """Process inline formatting."""
    parts = re.split(r'(\*\*.*?\*\*|__.*?__)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_chinese_font(run, bold=True)
        elif part.startswith('__') and part.endswith('__'):
            run = paragraph.add_run(part[2:-2])
            set_chinese_font(run, bold=True)
        else:
            if part:
                run = paragraph.add_run(part)
                set_chinese_font(run)


def find_image_file(image_dir, placeholder_index):
    """Find image file for a placeholder."""
    if not image_dir or not os.path.exists(image_dir):
        return None
    
    import glob
    
    # Support multiple naming patterns
    patterns = [
        f"img{placeholder_index:02d}_01.jpg",    # img01_01.jpg
        f"img{placeholder_index:02d}_01.png",
        f"img{placeholder_index:02d}_01.jpeg",
        f"img_{placeholder_index:02d}_01.jpg",  # img_01_01.jpg
        f"img_{placeholder_index:02d}_01.png",
        f"img_{placeholder_index:02d}_01.jpeg",
        f"img{placeholder_index:02d}.jpg",      # img01.jpg
        f"img_{placeholder_index:02d}.jpg",      # img_01.jpg
    ]
    
    for pattern in patterns:
        matches = glob.glob(os.path.join(image_dir, pattern))
        if matches:
            return matches[0]
    
    # Fallback: get all images and pick by index
    all_images = glob.glob(os.path.join(image_dir, "*.jpg")) + \
                 glob.glob(os.path.join(image_dir, "*.png")) + \
                 glob.glob(os.path.join(image_dir, "*.jpeg"))
    
    all_images.sort()  # Sort for consistent ordering
    
    if all_images and placeholder_index <= len(all_images):
        return all_images[placeholder_index - 1]
    
    return None


def add_image_to_doc(doc, image_path, description=""):
    """Add an image to the document with rounded corners."""
    try:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run()
        picture = run.add_picture(image_path, width=Inches(5.0))
        
        # Apply rounded corners to image
        try:
            from docx.oxml import parse_xml
            
            # Access the shape properties
            spPr = picture._inline.graphic.graphicData.pic.spPr
            
            # Create rounded rectangle geometry XML
            rounded_geom = parse_xml(
                r'<a:prstGeom xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" prst="roundRect">'
                r'<a:avLst>'
                r'<a:gd name="adj" fmla="val 5000"/>'
                r'</a:avLst>'
                r'</a:prstGeom>'
            )
            
            # Find and replace the existing geometry element
            # spPr is CT_ShapeProperties which has xPr element
            for child in spPr.getchildren():
                if child.tag.endswith('}xfrm'):
                    # Keep the transform, but we need to work with the geometry
                    pass
            
            # Alternative: Add preset geometry if not exists
            # Check if there's already a prstGeom or custGeom
            has_geom = False
            for child in spPr.getchildren():
                if 'Geom' in child.tag:
                    has_geom = True
                    # Replace it
                    spPr.remove(child)
                    break
            
            if not has_geom:
                # Insert after xfrm if exists
                xfrm_found = False
                for i, child in enumerate(spPr.getchildren()):
                    if child.tag.endswith('}xfrm'):
                        xfrm_found = True
                        # Insert after xfrm
                        spPr.insert(i + 1, rounded_geom)
                        break
                if not xfrm_found:
                    spPr.append(rounded_geom)
            else:
                # Re-add the rounded geometry
                xfrm_found = False
                for i, child in enumerate(spPr.getchildren()):
                    if child.tag.endswith('}xfrm'):
                        xfrm_found = True
                        spPr.insert(i + 1, rounded_geom)
                        break
                if not xfrm_found:
                    spPr.append(rounded_geom)
                    
        except Exception as e:
            # If rounding fails, just continue without it
            print(f"  Note: Could not apply rounded corners: {e}", file=sys.stderr)
        
        return True
    except Exception as e:
        print(f"  Warning: Could not insert image {image_path}: {e}", file=sys.stderr)
        return False


def add_table_to_doc(doc, headers, rows):
    """Add a table to the document."""
    try:
        # Create table
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    set_chinese_font(run, bold=True)
        
        # Add data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = cell_text
                    # Set font for all runs
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            set_chinese_font(run)
        
        # Add spacing after table
        doc.add_paragraph()
        
        return True
    except Exception as e:
        print(f"  Warning: Could not create table: {e}", file=sys.stderr)
        return False


def create_word_document(elements, output_path, image_dir=None, title=None):
    """Create Word document from parsed elements."""
    doc = Document()
    
    # Set default font to Microsoft YaHei
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    image_index = 0
    images_inserted = 0
    images_failed = 0
    tables_created = 0
    
    for element in elements:
        elem_type = element['type']
        
        if elem_type == 'heading1':
            para = doc.add_heading(element['content'], level=1)
            for run in para.runs:
                set_chinese_font(run, font_size=18, bold=True)
        
        elif elem_type == 'heading2':
            # Add empty paragraph before heading
            doc.add_paragraph()
            
            # Create heading with 1 space before and after content (total black background)
            content = f" {element['content']} "
            para = doc.add_heading(content, level=2)
            
            # Center align the heading
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Style: black background (including spaces), white font, 1.2x font size (14*1.2=16.8), italic
            for run in para.runs:
                set_chinese_font(run, font_size=17, bold=True)
                run.font.color.rgb = RGBColor(255, 255, 255)  # White font
                run.font.highlight_color = 1  # Black background (wdBlack)
                run.font.italic = True  # Italic style
            
            # Add empty paragraph after heading
            doc.add_paragraph()
        
        elif elem_type == 'heading3':
            # Heading 3: bold, slightly larger than body text, black color
            para = doc.add_heading(element['content'], level=3)
            for run in para.runs:
                set_chinese_font(run, font_size=14, bold=True)
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black font color
        
        elif elem_type == 'paragraph':
            para = doc.add_paragraph()
            process_inline_formatting(para, element['content'])
        
        elif elem_type == 'list':
            for item in element['items']:
                para = doc.add_paragraph(style='List Bullet')
                process_inline_formatting(para, item)
        
        elif elem_type == 'table':
            if add_table_to_doc(doc, element['headers'], element['rows']):
                tables_created += 1
        
        elif elem_type == 'image_placeholder':
            image_index += 1
            description = element['description']
            
            if image_dir:
                image_path = find_image_file(image_dir, image_index)
                if image_path and os.path.exists(image_path):
                    if add_image_to_doc(doc, image_path, description):
                        images_inserted += 1
                    else:
                        images_failed += 1
                else:
                    para = doc.add_paragraph()
                    run = para.add_run(f"[图片: {description}]")
                    set_chinese_font(run, font_size=10)
                    run.font.color.rgb = RGBColor(128, 128, 128)
                    run.italic = True
                    images_failed += 1
            else:
                para = doc.add_paragraph()
                run = para.add_run(f"[图片: {description}]")
                set_chinese_font(run, font_size=10)
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.italic = True
        
        elif elem_type == 'editor_line':
            # Editor/Reviewer lines: bold, smaller font, single line spacing
            para = doc.add_paragraph()
            para.paragraph_format.line_spacing = 1.0  # Single line spacing
            run = para.add_run(element['content'])
            set_chinese_font(run, font_size=10, bold=True)  # Smaller font (10pt) and bold
        
        elif elem_type == 'hr':
            pass  # Skip horizontal rules
    
    doc.save(output_path)
    
    return {
        'output_path': output_path,
        'images_inserted': images_inserted,
        'images_failed': images_failed,
        'tables_created': tables_created,
        'total_placeholders': image_index
    }


def convert_markdown_to_word(input_file, output_file=None, image_dir=None, title=None):
    """Convert Markdown file to Word document."""
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        if not output_file:
            output_file = input_file.rsplit('.', 1)[0] + '.docx'
        
        elements = parse_markdown(md_content)
        
        result = create_word_document(elements, output_file, image_dir, title)
        
        return {
            'success': True,
            'output_path': result['output_path'],
            'elements_count': len(elements),
            'images_inserted': result['images_inserted'],
            'images_failed': result['images_failed'],
            'tables_created': result['tables_created'],
            'total_placeholders': result['total_placeholders']
        }
    
    except Exception as e:
        return {
            'success': False,
            'error': str(e)
        }


def main():
    parser = argparse.ArgumentParser(description='Convert Markdown to Word document with images and tables')
    parser.add_argument('input', help='Input Markdown file path')
    parser.add_argument('-o', '--output', help='Output Word file path (optional)')
    parser.add_argument('-i', '--images', help='Directory containing images (optional)')
    parser.add_argument('-t', '--title', help='Document title (optional)')
    args = parser.parse_args()
    
    result = convert_markdown_to_word(args.input, args.output, args.images, args.title)
    
    if result['success']:
        print(f"✓ Successfully converted to: {result['output_path']}")
        print(f"✓ Processed {result['elements_count']} elements")
        if result['tables_created'] > 0:
            print(f"✓ Tables: {result['tables_created']} created")
        if result['total_placeholders'] > 0:
            print(f"✓ Images: {result['images_inserted']} inserted, {result['images_failed']} not found")
    else:
        print(f"✗ Error: {result['error']}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()